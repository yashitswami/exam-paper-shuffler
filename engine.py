import pandas as pd
import numpy as np
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_balanced_braces(text, start_pos):
    """
    Finds and returns the content of the balanced braces starting at start_pos (which points to '{').
    Returns (content, end_index)
    """
    if start_pos >= len(text) or text[start_pos] != '{':
        return None, start_pos
    
    depth = 0
    for i in range(start_pos, len(text)):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return text[start_pos+1:i], i + 1
    return None, start_pos

def replace_fractions(text):
    """
    Recursively replaces \frac{A}{B} and \dfrac{A}{B} with (A / B) using balanced brace matching.
    """
    for fn in [r'\dfrac', r'\frac']:
        while True:
            idx = text.find(fn)
            if idx == -1:
                break
            arg1_start = text.find('{', idx + len(fn))
            if arg1_start == -1:
                break
            num, arg1_end = extract_balanced_braces(text, arg1_start)
            if num is None:
                break
            
            arg2_start = text.find('{', arg1_end)
            if arg2_start == -1 or text[arg1_end:arg2_start].strip() != '':
                break
            den, arg2_end = extract_balanced_braces(text, arg2_start)
            if den is None:
                break
            
            replacement = f"({num.strip()} / {den.strip()})"
            text = text[:idx] + replacement + text[arg2_end:]
    return text

def replace_overunder_sets(text):
    """
    Replaces \overset{top}{base} and \underset{bottom}{base} with base(top) / base(bottom)
    for clean chemical structural representations.
    """
    for fn in [r'\overset', r'\underset']:
        while True:
            idx = text.find(fn)
            if idx == -1:
                break
            arg1_start = text.find('{', idx + len(fn))
            if arg1_start == -1:
                break
            top, arg1_end = extract_balanced_braces(text, arg1_start)
            if top is None:
                break
            
            arg2_start = text.find('{', arg1_end)
            if arg2_start == -1:
                break
            base, arg2_end = extract_balanced_braces(text, arg2_start)
            if base is None:
                break
            
            # Clean up structural bond symbols inside top/base
            top_clean = top.replace(r'\mid', '').replace('|', '').replace('=', '').replace(r'\textbardbl', '=').strip()
            top_clean = re.sub(r'^\{+|\}$', '', top_clean)
            base_clean = base.replace(r'\mid', '').replace('|', '').strip()
            base_clean = re.sub(r'^\{+|\}$', '', base_clean)
            
            if top_clean and base_clean and base_clean not in ['=', r'\textbardbl']:
                replacement = f"{base_clean}({top_clean})"
            elif top_clean:
                replacement = f"({top_clean})"
            else:
                replacement = base_clean
                
            text = text[:idx] + replacement + text[arg2_end:]
    return text

def sanitize_text_for_word(text):
    """
    Converts raw LaTeX equations, structural formulas, and units into clean, 
    human-readable plain text for Word documents.
    """
    if pd.isna(text) or text is None:
        return ""
    
    t = str(text).strip()
    
    # Preserve URLs before applying LaTeX transformations
    urls = []
    def preserve_url(match):
        urls.append(match.group(0).replace(r'\_', '_'))
        return f"__URL_PLACEHOLDER_{len(urls)-1}__"
    
    t = re.sub(r'https?://[^\s]+', preserve_url, t)
    
    # 1. Unescape HTML / URL entities
    t = t.replace('&lt;', '<').replace('&gt;', '>')
    t = t.replace(r'\lt', '<').replace(r'\gt', '>')
    
    # 2. Structural chemical Overset / Underset replacement
    t = replace_overunder_sets(t)

    # 3. Fractions replacement with balanced brace handling
    t = replace_fractions(t)

    # 4. Clean Matrices / Arrays / Tables
    t = re.sub(r'\\begin\{(matrix|array|align|equation|table)\}(?:\[[^\]]*\])?(?:\{[^}]*\})?', '', t)
    t = re.sub(r'\\end\{(matrix|array|align|equation|table)\}', '', t)
    t = t.replace(r'\\', ' ').replace(r'\hline', '').replace(r'\quad', ' ')
    t = re.sub(r'\s*&\s*', ' | ', t)
    
    # 5. Clean Text Formatting Commands
    t = re.sub(r'\\text(?:bf|it|rm|sf|tt)?\{([^}]+)\}', r'\1', t)
    t = re.sub(r'\\math(?:rm|bf|it|sf|bb|cal)?\{([^}]+)\}', r'\1', t)
    t = t.replace(r'\displaystyle', '')
    
    # 6. Vectors and Unit Vectors
    t = re.sub(r'\\(?:overrightarrow|vec)\{([^}]+)\}', r'\1', t)
    t = re.sub(r'\\(?:widehat|hat)\{([ijk])\}', r'\1̂', t)
    t = re.sub(r'\\(?:widehat|hat)\{([^}]+)\}', r'\1', t)
    
    # 7. Square Roots (\sqrt)
    t = re.sub(r'\\sqrt\[([^\]]+)\]\{([^}]+)\}', r'\1√(\2)', t)
    t = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', t)
    t = re.sub(r'\\sqrt\s*([a-zA-Z0-9]+)', r'√\1', t)
    
    # 8. Delimiters (\left, \right, \lbrack, \rbrack)
    t = re.sub(r'\\(left|right)\b', '', t)
    t = t.replace(r'\leftlbrack', '[').replace(r'\rightrbrack', ']')
    t = t.replace(r'\lbrack', '[').replace(r'\rbrack', ']')
    t = t.replace(r'\{', '{').replace(r'\}', '}')
    
    # 9. Inverse Trig & Math Functions (e.g., \sin^{-1}, \tan^{-1})
    t = re.sub(r'\\(sin|cos|tan|cot|sec|csc)\^\{\s*\-\s*1\s*\}', r'\1⁻¹', t)
    t = re.sub(r'\\(sin|cos|tan|cot|sec|csc)\b', r'\1', t)
    t = re.sub(r'\\(log|ln)\b', r'\1', t)
    
    # 10. Greek letters, Operators, & Symbols
    replacements = [
        (r'\\alpha\b', 'α'), (r'\\beta\b', 'β'), (r'\\gamma\b', 'γ'), (r'\\delta\b', 'δ'),
        (r'\\epsilon\b', 'ε'), (r'\\theta\b', 'θ'), (r'\\lambda\b', 'λ'), (r'\\mu\b', 'μ'),
        (r'\\pi\b', 'π'), (r'\\sigma\b', 'σ'), (r'\\omega\b', 'ω'), (r'\\Delta\b', 'Δ'),
        (r'\\Omega\b', 'Ω'), (r'\\times\b', '×'), (r'\\div\b', '÷'), (r'\\pm\b', '±'),
        (r'\\leq?\b', '≤'), (r'\\geq?\b', '≥'), (r'\\neq\b', '≠'), (r'\\approx\b', '≈'),
        (r'\\infty\b', '∞'), (r'\\rightarrow\b', '→'), (r'\\leftarrow\b', '←'),
        (r'\\Rightarrow\b', '⇒'), (r'\\degree\b', '°'), (r'\\circ\b', '°'),
        (r'\\textemdash\b', '-'), (r'\\cdot\b', '·'), (r'\\textbardbl\b', '='), (r'\\colon\b', ':')
    ]
    for pattern, symbol in replacements:
        t = re.sub(pattern, symbol, t)

    # Clean Degree Notation artifacts like {^{°}}, {^°}, {°}
    t = re.sub(r'\{\s*\^?\s*°\s*\}', '°', t)
    t = re.sub(r'\^\s*°', '°', t)
    t = t.replace('{°}', '°')
        
    # 11. Superscripts and Subscripts (handles negative signs and spaces gracefully)
    sub_map = str.maketrans("0123456789+-=() ", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ ")
    sup_map = str.maketrans("0123456789+-=() ", "⁰¹²³⁴⁵₆₇⁸⁹⁺⁻⁼⁽⁾ ")
    
    # Translate _{...} and ^{...}
    t = re.sub(r'\_\{([0-9+\-=()\s]+)\}', lambda m: m.group(1).translate(sub_map), t)
    t = re.sub(r'\^\{([0-9+\-=()\s]+)\}', lambda m: m.group(1).translate(sup_map), t)
    t = re.sub(r'\_([0-9])', lambda m: m.group(1).translate(sub_map), t)
    t = re.sub(r'\^([0-9])', lambda m: m.group(1).translate(sup_map), t)
    
    # Normalize common units / exponents
    t = t.replace('⁻ 1', '⁻¹').replace('⁻ 2', '⁻²').replace('⁻ 3', '⁻³')
    t = t.replace('⁺ 1', '⁺¹').replace('⁺ 2', '⁺²')
    t = t.replace('^-1', '⁻¹').replace('^-2', '⁻²').replace('^-3', '⁻³')
    t = t.replace('^2', '²').replace('^3', '³')

    # 12. Final Cleanup
    t = t.replace('$', '').replace(r'\(', '').replace(r'\)', '')
    t = re.sub(r'\\([a-zA-Z]+)', r'\1', t)
    t = re.sub(r'\\\s*', ' ', t)
    
    # Restore URLs
    for idx, url in enumerate(urls):
        t = t.replace(f"__URL_PLACEHOLDER_{idx}__", url)
        
    # Clean up double nested parentheses e.g. ( (4R / H) ) -> (4R / H)
    t = re.sub(r'\(\s*\(\s*([^()]+)\s*\)\s*\)', r'(\1)', t)
    
    # Remove artificial outer bracket wraps around numbers/units e.g. (60 kg) -> 60 kg
    t = re.sub(r'(?<![a-zA-Z0-9/])\(\s*([a-zA-Z0-9\s.=+\-²³α-ωΑ-Ω]+)\s*\)(?![a-zA-Z0-9/])', r'\1', t)

    t = re.sub(r'[ \t]+', ' ', t)
    
    return t.strip()

def process_and_shuffle(excel_file, num_sets, exam_title, time_limit, max_marks, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    
    # Load Excel Data
    df = pd.read_excel(excel_file)
    df['Subject'] = df['Unnamed: 0'].ffill()
    
    sets = [f"Set {chr(65+i)}" for i in range(num_sets)]
    
    # Subject section boundaries (JEE Main standard layout)
    sec_ranges = [
        ('Physics', 1, 20, 21, 25),
        ('Chemistry', 26, 45, 46, 50),
        ('Mathematics', 51, 70, 71, 75)
    ]
    
    # Perform Shuffling Matrix Generation
    set_mappings = {}
    for s in sets:
        set_q_map = {}
        for subj, scq_start, scq_end, int_start, int_end in sec_ranges:
            # 1. Shuffle Single Choice Questions (SCQs)
            scq_shuffled = list(range(scq_start, scq_end + 1))
            np.random.shuffle(scq_shuffled)
            for idx, master_q in enumerate(scq_shuffled):
                opts = ['A', 'B', 'C', 'D']
                np.random.shuffle(opts)
                opt_map_str = f"A->{chr(65+opts.index('A'))}, B->{chr(65+opts.index('B'))}, C->{chr(65+opts.index('C'))}, D->{chr(65+opts.index('D'))}"
                set_q_map[master_q] = {
                    'new_q': scq_start + idx,
                    'opt_order': opts,
                    'opt_map_str': opt_map_str
                }
                
            # 2. Shuffle Numerical / Integer Type Questions
            int_shuffled = list(range(int_start, int_end + 1))
            np.random.shuffle(int_shuffled)
            for idx, master_q in enumerate(int_shuffled):
                set_q_map[master_q] = {
                    'new_q': int_start + idx,
                    'opt_order': ['A', 'B', 'C', 'D'],
                    'opt_map_str': 'N/A'
                }
        set_mappings[s] = set_q_map

    # 1. Create Master Shuffling Matrix CSV for EvalBee
    mapping_rows = []
    total_questions = len(df)
    for q in range(1, total_questions + 1):
        subj = 'Physics' if q <= 25 else ('Chemistry' if q <= 50 else 'Mathematics')
        q_type = 'SCQ' if (q <= 20 or (26 <= q <= 45) or (51 <= q <= 70)) else 'Integer'
        row = {'Master_Q': q, 'Subject': subj, 'Type': q_type}
        for s in sets:
            m = set_mappings[s][q]
            row[f'{s}_QNo'] = m['new_q']
            row[f'{s}_OptMap'] = m['opt_map_str']
        mapping_rows.append(row)
        
    csv_path = os.path.join(output_dir, "Master_Shuffling_Matrix.csv")
    pd.DataFrame(mapping_rows).to_csv(csv_path, index=False)
    
    # 2. Build Word Documents (.docx) for Each Set
    generated_files = [csv_path]
    for s in sets:
        doc = Document()
        
        # Configure page margins (0.7 inches)
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
            
        # Header Section
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r1 = header_p.add_run(f"{exam_title}\n")
        r1.bold = True
        r1.font.size = Pt(16)
        
        r2 = header_p.add_run(f"TEST PAPER: {s}\n")
        r2.bold = True
        r2.font.size = Pt(14)
        r2.font.color.rgb = RGBColor(0, 51, 102)
        
        meta_p = doc.add_paragraph(f"Time: {time_limit}\t\t\t\t\tMaximum Marks: {max_marks}")
        meta_p.runs[0].bold = True
        meta_p.runs[0].font.size = Pt(10.5)
        
        doc.add_paragraph("-" * 80)
        
        # Invert mapping to iterate sequentially by new question number (Q1 to Q75)
        new_to_master = {info['new_q']: (mq, info) for mq, info in set_mappings[s].items()}
        current_subj = None
        
        for new_q in range(1, total_questions + 1):
            mq, info = new_to_master[new_q]
            row = df[df['Q.No'] == mq].iloc[0]
            subj = row['Subject']
            
            # Subject Heading
            if subj != current_subj:
                current_subj = subj
                subj_p = doc.add_paragraph()
                sr = subj_p.add_run(f"\nSECTION: {current_subj.upper()}")
                sr.bold = True
                sr.font.size = Pt(13)
                sr.font.color.rgb = RGBColor(153, 0, 0)
                doc.add_paragraph("=" * 60)
                
            # Integer Sub-section Heading
            if new_q in [21, 46, 71]:
                num_p = doc.add_paragraph()
                nr = num_p.add_run("Sub-Section: Numerical / Integer Type (Answer as Numerical Value)")
                nr.bold = True
                nr.italic = True
                nr.font.size = Pt(11)
                
            # Question Text
            qp = doc.add_paragraph()
            q_num_run = qp.add_run(f"Q.{new_q} ")
            q_num_run.bold = True
            q_num_run.font.size = Pt(10.5)
            
            q_text_run = qp.add_run(sanitize_text_for_word(row['Qus']))
            q_text_run.font.size = Pt(10.5)
            
            # SCQ Options vs Numerical Space
            is_integer_q = new_q in [21, 22, 23, 24, 25, 46, 47, 48, 49, 50, 71, 72, 73, 74, 75] or pd.isna(row['Opt A'])
            
            if not is_integer_q:
                tbl = doc.add_table(rows=2, cols=2)
                coords = [(0, 0), (0, 1), (1, 0), (1, 1)]
                labels = ['(a)', '(b)', '(c)', '(d)']
                
                for idx, orig_letter in enumerate(info['opt_order']):
                    r, c = coords[idx]
                    cell = tbl.cell(r, c)
                    cell.width = Inches(3.4)
                    
                    cell_p = cell.paragraphs[0]
                    cell_p.paragraph_format.space_before = Pt(2)
                    cell_p.paragraph_format.space_after = Pt(2)
                    
                    lbl_run = cell_p.add_run(f"{labels[idx]} ")
                    lbl_run.bold = True
                    lbl_run.font.size = Pt(10)
                    
                    opt_val = sanitize_text_for_word(row[f'Opt {orig_letter}'])
                    val_run = cell_p.add_run(opt_val)
                    val_run.font.size = Pt(10)
                    
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            else:
                space_p = doc.add_paragraph()
                space_p.paragraph_format.space_after = Pt(10)
                
        docx_file = os.path.join(output_dir, f"Question_Paper_{s.replace(' ', '_')}.docx")
        doc.save(docx_file)
        generated_files.append(docx_file)
        
    return generated_files
